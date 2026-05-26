"""
Genera Scoring_Models.docx — documentación visual de scoring/models.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────── palette ────────────────────────────────────────

P = {
    "dark_blue":    (0x0F, 0x3E, 0x5A),
    "green":        (0x2D, 0x9B, 0x4E),
    "teal":         (0x0E, 0x76, 0x8C),
    "amber":        (0xB4, 0x5A, 0x00),
    "purple":       (0x5B, 0x2D, 0x8C),
    "red":          (0x9B, 0x1B, 0x1B),
    "white":        (0xFF, 0xFF, 0xFF),
    "light_gray":   (0xF5, 0xF5, 0xF5),
    "mid_gray":     (0x88, 0x88, 0x88),
    "dark_text":    (0x22, 0x22, 0x22),
}

HEX = {
    "dark_blue":    "0F3E5A",
    "green":        "2D9B4E",
    "teal":         "0E768C",
    "amber":        "B45A00",
    "purple":       "5B2D8C",
    "red":          "9B1B1B",
    "white":        "FFFFFF",
    "light_blue":   "EAF4FB",
    "light_green":  "F0FFF4",
    "light_amber":  "FFF8F0",
    "light_purple": "F5F0FF",
    "light_red":    "FFF0F0",
    "light_teal":   "F0FAFA",
    "stripe_a":     "F5FBFF",
    "stripe_b":     "FFFFFF",
    "section_bg":   "E8F4FA",
}

# color per model
MODEL_COLOR = {
    "ScoringSession":          ("teal",   "light_teal"),
    "ScoringSessionPlayer":    ("green",  "light_green"),
    "ScoringEvent":            ("amber",  "light_amber"),
    "MaterializedScore":       ("purple", "light_purple"),
    "LeaderboardEntry":        ("red",    "light_red"),
    "ScoringMediaAttachment":  ("dark_blue", "light_blue"),
}

# ─────────────────────────── xml helpers ────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color_hex, size="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_row_height(row, pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(pt * 20)))
    trPr.append(trHeight)

def make_run(para, text, bold=False, italic=False, size_pt=9, color=None, font=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if font:
        r.font.name = font
    return r

# ─────────────────────────── layout helpers ─────────────────────────────────

def section_divider(doc, title, color_key="teal"):
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, HEX[color_key])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"  {title}")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(*P["white"])
    doc.add_paragraph()

def field_table(doc, rows, header_color, row_bg_a, row_bg_b):
    """
    rows = list of dicts with keys: name, type, nullable, default, description, pk, fk, unique
    """
    COL_W = [Inches(1.5), Inches(1.1), Inches(0.55), Inches(0.55), Inches(3.4)]
    HEADERS = ["Campo", "Tipo / Choices", "Null", "Default", "Descripción / Notas"]

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = tbl.rows[0]
    set_row_height(hdr, 14)
    for i, (txt, w) in enumerate(zip(HEADERS, COL_W)):
        c = hdr.cells[i]
        c.width = w
        set_cell_bg(c, header_color)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(*P["white"])

    # data rows
    for idx, row in enumerate(rows):
        tr = tbl.add_row()
        set_row_height(tr, 13)
        bg = row_bg_a if idx % 2 == 0 else row_bg_b

        for ci in range(5):
            tr.cells[ci].width = COL_W[ci]
            set_cell_bg(tr.cells[ci], bg)

        # col 0 — field name
        c0 = tr.cells[0]
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(1)
        p0.paragraph_format.space_after = Pt(1)
        name_txt = row["name"]
        prefixes = []
        if row.get("pk"):    prefixes.append("🔑")
        if row.get("fk"):    prefixes.append("→")
        if row.get("unique"): prefixes.append("◆")
        prefix = " ".join(prefixes) + " " if prefixes else ""
        r0 = p0.add_run(prefix + name_txt)
        r0.bold = True
        r0.font.size = Pt(8)
        r0.font.name = "Consolas"
        r0.font.color.rgb = RGBColor(*P["dark_text"])

        # col 1 — type
        p1 = tr.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(1)
        r1 = p1.add_run(row.get("type", ""))
        r1.font.size = Pt(8)
        r1.font.name = "Consolas"
        r1.font.color.rgb = RGBColor(0x11, 0x55, 0x88)

        # col 2 — nullable
        p2 = tr.cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        null_val = row.get("nullable", "")
        r2 = p2.add_run(null_val)
        r2.font.size = Pt(8)
        if null_val == "Sí":
            r2.font.color.rgb = RGBColor(0xAA, 0x66, 0x00)
        else:
            r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # col 3 — default
        p3 = tr.cells[3].paragraphs[0]
        p3.paragraph_format.space_before = Pt(1)
        r3 = p3.add_run(row.get("default", "—"))
        r3.font.size = Pt(7.5)
        r3.font.name = "Consolas"
        r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # col 4 — description
        p4 = tr.cells[4].paragraphs[0]
        p4.paragraph_format.space_before = Pt(1)
        r4 = p4.add_run(row.get("description", ""))
        r4.font.size = Pt(8)
        r4.font.color.rgb = RGBColor(*P["dark_text"])

    doc.add_paragraph()

def enum_box(doc, name, values_desc, color_key):
    """Small table showing an enum's values."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    make_run(p, f"  Enum: {name}", bold=True, size_pt=9, color=P[color_key])

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0]
    for i, txt in enumerate(["Valor", "Descripción"]):
        c = hdr.cells[i]
        set_cell_bg(c, HEX[color_key])
        p2 = c.paragraphs[0]
        r = p2.add_run(txt)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(*P["white"])

    for idx, (val, desc) in enumerate(values_desc):
        tr = tbl.add_row()
        bg = "F8F8F8" if idx % 2 == 0 else "FFFFFF"
        set_cell_bg(tr.cells[0], bg)
        set_cell_bg(tr.cells[1], bg)

        p_v = tr.cells[0].paragraphs[0]
        rv = p_v.add_run(val)
        rv.font.size = Pt(8)
        rv.font.name = "Consolas"
        rv.font.color.rgb = RGBColor(0x11, 0x55, 0x88)

        p_d = tr.cells[1].paragraphs[0]
        rd = p_d.add_run(desc)
        rd.font.size = Pt(8)

    for row in tbl.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(5.3)

    doc.add_paragraph()

def model_header(doc, model_name, description, mobile_equiv):
    doc.add_page_break()
    color_key, bg_key = MODEL_COLOR[model_name]

    # big colored banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, HEX[color_key])
    cell.width = Inches(7.1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    make_run(p, f"  {model_name}", bold=True, size_pt=15, color=P["white"])

    doc.add_paragraph()

    # description
    if description:
        pd = doc.add_paragraph()
        pd.paragraph_format.space_after = Pt(4)
        make_run(pd, description, size_pt=9, color=P["dark_text"])

    # mobile equivalent
    if mobile_equiv:
        pm = doc.add_paragraph()
        pm.paragraph_format.space_after = Pt(8)
        make_run(pm, "Equivalente en mobile: ", bold=True, size_pt=9, color=P[color_key])
        make_run(pm, mobile_equiv, italic=True, size_pt=9, color=P["mid_gray"])

    doc.add_paragraph()

def rel_table(doc, rels):
    """rels = list of (field, target_model, on_delete, note)"""
    p = doc.add_paragraph()
    make_run(p, "  Relaciones (Foreign Keys)", bold=True, size_pt=9, color=P["dark_blue"])

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0]
    for i, txt in enumerate(["Campo FK", "Modelo destino", "on_delete", "Nota"]):
        set_cell_bg(hdr.cells[i], HEX["dark_blue"])
        r = hdr.cells[i].paragraphs[0].add_run(txt)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(*P["white"])

    for idx, (field, target, on_del, note) in enumerate(rels):
        tr = tbl.add_row()
        bg = "F5FBFF" if idx % 2 == 0 else "FFFFFF"
        for c in tr.cells:
            set_cell_bg(c, bg)

        def rc(ci, text, bold=False, mono=False, clr=None):
            p2 = tr.cells[ci].paragraphs[0]
            r2 = p2.add_run(text)
            r2.bold = bold
            r2.font.size = Pt(8)
            if mono:
                r2.font.name = "Consolas"
            if clr:
                r2.font.color.rgb = RGBColor(*clr)

        rc(0, field, bold=True, mono=True)
        rc(1, target, mono=True, clr=P["teal"])
        rc(2, on_del, mono=True, clr=(0xAA, 0x44, 0x00))
        rc(3, note)

    for row in tbl.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(1.7)
        row.cells[2].width = Inches(0.9)
        row.cells[3].width = Inches(3.0)

    doc.add_paragraph()

def index_table(doc, indexes):
    p = doc.add_paragraph()
    make_run(p, "  Índices de base de datos", bold=True, size_pt=9, color=P["dark_blue"])

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0]
    for i, txt in enumerate(["Tipo / Constraint", "Campos"]):
        set_cell_bg(hdr.cells[i], HEX["dark_blue"])
        r = hdr.cells[i].paragraphs[0].add_run(txt)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(*P["white"])

    for idx, (tipo, campos) in enumerate(indexes):
        tr = tbl.add_row()
        bg = "F5FBFF" if idx % 2 == 0 else "FFFFFF"
        set_cell_bg(tr.cells[0], bg)
        set_cell_bg(tr.cells[1], bg)
        r0 = tr.cells[0].paragraphs[0].add_run(tipo)
        r0.font.size = Pt(8)
        r0.bold = True
        r1 = tr.cells[1].paragraphs[0].add_run(campos)
        r1.font.size = Pt(8)
        r1.font.name = "Consolas"
        r1.font.color.rgb = RGBColor(0x11, 0x55, 0x88)

    for row in tbl.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(5.5)

    doc.add_paragraph()

# ─────────────────────────── MODEL DATA ─────────────────────────────────────

BASE_FIELDS = [
    {"name": "id",         "type": "AutoField (PK)",  "nullable": "No",  "default": "auto",  "description": "Clave primaria entera autogenerada por Django.", "pk": True},
    {"name": "uuid",       "type": "UUIDField",       "nullable": "No",  "default": "uuid4", "description": "UUID único expuesto al exterior. Inmutable tras la creación.", "unique": True},
    {"name": "created_at", "type": "DateTimeField",   "nullable": "No",  "default": "now",   "description": "Fecha y hora de creación. Rellenado automáticamente."},
    {"name": "updated_at", "type": "DateTimeField",   "nullable": "No",  "default": "now",   "description": "Fecha y hora de última modificación. Actualizado automáticamente."},
]

MODELS = [

    # ── ScoringSession ────────────────────────────────────────────────────
    {
        "name": "ScoringSession",
        "description": (
            "Representa una ronda completa puntuada por un grupo de jugadores. "
            "Una sesión = un dispositivo actuando como scorer para un grupo en una ronda. "
            "Los jugadores se asocian mediante ScoringSessionPlayer; los scores individuales "
            "viven en MaterializedScore; el historial inmutable de eventos en ScoringEvent."
        ),
        "mobile_equiv": "rounds  (tabla local WatermelonDB)",
        "fields": BASE_FIELDS + [
            {"name": "tour_event",               "type": "FK → TourEvent",         "nullable": "Sí", "default": "NULL",  "description": "Evento de competición. NULL cuando mode = free-play.", "fk": True},
            {"name": "tour_event_start_time",     "type": "FK → TourEventStartTime","nullable": "Sí", "default": "NULL",  "description": "Slot de salida (tee-time) que puntúa esta sesión. Resuelve codigo_grupo → session.", "fk": True},
            {"name": "course",                   "type": "FK → Course",            "nullable": "Sí", "default": "NULL",  "description": "Campo de golf donde se juega la ronda.", "fk": True},
            {"name": "route",                    "type": "FK → Route",             "nullable": "Sí", "default": "NULL",  "description": "Variante de scorecard (por tee). Desambigua qué filas Hole usa el replay.", "fk": True},
            {"name": "mode",                     "type": "CharField(20)",           "nullable": "No", "default": "competition", "description": "Modo de la sesión. Choices: ScoringSessionMode."},
            {"name": "status",                   "type": "CharField(20)",           "nullable": "No", "default": "in_progress", "description": "Ciclo de vida de la sesión. Choices: ScoringSessionStatus."},
            {"name": "started_at",               "type": "DateTimeField",           "nullable": "Sí", "default": "NULL",  "description": "Momento en que el primer jugador marcó PLAYER_READY."},
            {"name": "completed_at",             "type": "DateTimeField",           "nullable": "Sí", "default": "NULL",  "description": "Momento en que la ronda terminó (último hoyo guardado)."},
            {"name": "locked_at",                "type": "DateTimeField",           "nullable": "Sí", "default": "NULL",  "description": "Se establece cuando la tarjeta se firma/finaliza. Sin enmiendas sin admin."},
            {"name": "tee_color",                "type": "CharField(20)",           "nullable": "Sí", "default": "NULL",  "description": "Color de tee de la sesión completa. Choices: TeeColor."},
            {"name": "group_code",               "type": "CharField(32)",           "nullable": "Sí", "default": "NULL",  "description": "Vincula dispositivos que puntúan el mismo grupo. Equivale a codigo_grupo en el móvil."},
            {"name": "hole_pars_snapshot",       "type": "JSONField",               "nullable": "No", "default": "{}",    "description": "Snapshot inmutable de los pares de hoyo al inicio de la ronda."},
            {"name": "hole_handicaps_snapshot",  "type": "JSONField",               "nullable": "No", "default": "{}",    "description": "Snapshot inmutable de los índices de handicap de hoyo al inicio de la ronda."},
            {"name": "meta",                     "type": "JSONField",               "nullable": "No", "default": "{}",    "description": "Metadatos libres (ej. nombre_competicion, nombre_prueba para free-play)."},
        ],
        "relations": [
            ("tour_event",            "tour.TourEvent",            "SET_NULL",  "NULL en partidas libres."),
            ("tour_event_start_time", "tour.TourEventStartTime",   "SET_NULL",  "Resuelve el grupo_code al crear la sesión."),
            ("course",                "club.Course",               "PROTECT",   "El campo no se puede borrar si tiene sesiones."),
            ("route",                 "club.Route",                "PROTECT",   "El recorrido no se puede borrar si tiene sesiones."),
        ],
        "indexes": [
            ("Index",        "tour_event, status"),
            ("Index",        "status, -started_at"),
            ("ordering",     "-created_at"),
        ],
    },

    # ── ScoringSessionPlayer ──────────────────────────────────────────────
    {
        "name": "ScoringSessionPlayer",
        "description": (
            "Jugadores participantes en una sesión de puntuación. "
            "Almacena el handicap y el tee en el momento de la ronda, "
            "el estado del jugador y el dispositivo vinculado."
        ),
        "mobile_equiv": "round_players  (tabla local WatermelonDB)",
        "fields": BASE_FIELDS + [
            {"name": "session",             "type": "FK → ScoringSession",       "nullable": "No", "default": "—",    "description": "Sesión a la que pertenece el jugador.", "fk": True},
            {"name": "player",              "type": "FK → core.User",            "nullable": "No", "default": "—",    "description": "Usuario que participa en la sesión.", "fk": True},
            {"name": "registration",        "type": "FK → TourEventRegistration","nullable": "Sí", "default": "NULL", "description": "Inscripción al evento cuando mode = competition.", "fk": True},
            {"name": "tee_color",           "type": "CharField(20)",              "nullable": "Sí", "default": "NULL", "description": "Color de tee asignado al jugador. Choices: TeeColor."},
            {"name": "handicap_index",      "type": "DecimalField(4,1)",          "nullable": "Sí", "default": "NULL", "description": "Snapshot del handicap index del jugador al inicio de la ronda."},
            {"name": "playing_handicap",    "type": "PositiveSmallIntegerField",  "nullable": "Sí", "default": "NULL", "description": "Playing handicap ajustado al campo para el cálculo del neto. Entero."},
            {"name": "status",              "type": "CharField(20)",              "nullable": "No", "default": "not_started", "description": "Estado del jugador en la sesión. Choices: ScoringSessionPlayerStatus."},
            {"name": "linked_device_hash",  "type": "CharField(64)",              "nullable": "Sí", "default": "NULL", "description": "Hash del X-Device-ID del dispositivo que reclamó este slot vía /link-device/."},
        ],
        "relations": [
            ("session",       "scoring.ScoringSession",        "CASCADE",   "Eliminar la sesión borra sus jugadores."),
            ("player",        "core.User",                     "CASCADE",   "Eliminar el usuario borra sus participaciones."),
            ("registration",  "tour.TourEventRegistration",    "SET_NULL",  "La inscripción puede borrarse sin afectar el score."),
        ],
        "indexes": [
            ("unique_together", "session, player"),
        ],
        "status_transitions": True,
    },

    # ── ScoringEvent ──────────────────────────────────────────────────────
    {
        "name": "ScoringEvent",
        "description": (
            "Log append-only de eventos de puntuación (equivale a action_log en el móvil). "
            "event_id es el UUID v7 generado por el cliente y actúa como clave de idempotencia "
            "para POST /api/v1/sync/ — reenviar el mismo id es un no-op. "
            "uuid (de BaseModel) es el identificador del lado servidor."
        ),
        "mobile_equiv": "action_log  (tabla local WatermelonDB)",
        "fields": BASE_FIELDS + [
            {"name": "session",                "type": "FK → ScoringSession",  "nullable": "No", "default": "—",       "description": "Sesión a la que pertenece el evento.", "fk": True},
            {"name": "event_id",               "type": "UUIDField",             "nullable": "No", "default": "—",       "description": "UUID v7 generado por el cliente. Clave de idempotencia para sync.", "unique": True},
            {"name": "action_type",            "type": "CharField(32)",          "nullable": "No", "default": "—",       "description": "Tipo de acción. Choices: ScoringActionType."},
            {"name": "device_id",              "type": "CharField(128)",         "nullable": "Sí", "default": "NULL",    "description": "Dispositivo que produjo el evento (de expo-secure-store)."},
            {"name": "sequence",               "type": "PositiveIntegerField",   "nullable": "Sí", "default": "NULL",    "description": "Contador monotónico por (session, device). Define el orden causal."},
            {"name": "client_timestamp",       "type": "DateTimeField",          "nullable": "No", "default": "—",       "description": "Timestamp del cliente cuando ocurrió la acción."},
            {"name": "server_received_at",     "type": "DateTimeField",          "nullable": "No", "default": "auto_now_add", "description": "Timestamp de anclaje — momento en que el servidor aceptó el evento."},
            {"name": "payload",                "type": "JSONField",              "nullable": "No", "default": "{}",      "description": "Datos específicos del tipo de acción (hole_number, strokes, player_id, etc.)."},
            {"name": "materialization_status", "type": "CharField(20)",          "nullable": "No", "default": "pending", "description": "Estado de materialización del evento. Choices: ScoringMaterializationStatus."},
            {"name": "materialization_error",  "type": "TextField",              "nullable": "Sí", "default": "NULL",    "description": "Mensaje de error si la materialización falló."},
        ],
        "relations": [
            ("session", "scoring.ScoringSession", "CASCADE", "Eliminar la sesión borra todos sus eventos."),
        ],
        "indexes": [
            ("Index",    "session, sequence"),
            ("Index",    "session, action_type"),
            ("Index",    "session, materialization_status"),
            ("unique",   "event_id"),
            ("ordering", "session, sequence, server_received_at"),
        ],
    },

    # ── MaterializedScore ─────────────────────────────────────────────────
    {
        "name": "MaterializedScore",
        "description": (
            "Score materializado por hoyo y jugador (equivale a hole_scores en el móvil). "
            "Proyección optimizada para lectura, reconstruida a partir del replay de ScoringEvent. "
            "El event log sigue siendo la fuente de verdad — esta tabla se puede reconstruir siempre."
        ),
        "mobile_equiv": "hole_scores  (tabla local WatermelonDB)",
        "fields": BASE_FIELDS + [
            {"name": "session",              "type": "FK → ScoringSession", "nullable": "No", "default": "—",    "description": "Sesión de puntuación a la que pertenece el score.", "fk": True},
            {"name": "hole",                 "type": "FK → club.Hole",     "nullable": "No", "default": "—",    "description": "Hoyo al que corresponde el score.", "fk": True},
            {"name": "player",               "type": "FK → core.User",     "nullable": "No", "default": "—",    "description": "Jugador al que corresponde el score.", "fk": True},
            {"name": "strokes",              "type": "PositiveSmallIntegerField", "nullable": "Sí", "default": "NULL", "description": "Golpes brutos registrados para ese hoyo."},
            {"name": "putts",                "type": "PositiveSmallIntegerField", "nullable": "Sí", "default": "NULL", "description": "Número de putts."},
            {"name": "penalties",            "type": "PositiveSmallIntegerField", "nullable": "No", "default": "0",    "description": "Golpes de penalización."},
            {"name": "fairway_hit",          "type": "BooleanField",              "nullable": "Sí", "default": "NULL", "description": "True si el jugador alcanzó el fairway en el tee shot."},
            {"name": "gir",                  "type": "BooleanField",              "nullable": "Sí", "default": "NULL", "description": "Green in Regulation: True si el jugador llegó al green en la regulación."},
            {"name": "points",               "type": "IntegerField",              "nullable": "Sí", "default": "NULL", "description": "Puntos Stableford calculados, si aplica."},
            {"name": "net_strokes",          "type": "IntegerField",              "nullable": "Sí", "default": "NULL", "description": "Golpes netos (strokes − handicap stroke en el hoyo)."},
            {"name": "last_event",           "type": "FK → ScoringEvent",        "nullable": "Sí", "default": "NULL", "description": "Evento cuyo replay produjo los valores actuales.", "fk": True},
            {"name": "flagged_conflict",     "type": "BooleanField",              "nullable": "No", "default": "False", "description": "True si el pipeline detectó dos eventos concurrentes (< 5s) en la misma sesión."},
            {"name": "conflict_detected_at", "type": "DateTimeField",             "nullable": "Sí", "default": "NULL", "description": "Timestamp en que flagged_conflict se puso a True por última vez."},
        ],
        "relations": [
            ("session",    "scoring.ScoringSession", "CASCADE",   "Eliminar la sesión borra todos sus scores."),
            ("hole",       "club.Hole",              "CASCADE",   "Eliminar el hoyo borra sus scores materializados."),
            ("player",     "core.User",              "CASCADE",   "Eliminar el usuario borra sus scores."),
            ("last_event", "scoring.ScoringEvent",   "SET_NULL",  "El evento de referencia puede borrarse; el score permanece."),
        ],
        "indexes": [
            ("unique_together", "session, hole, player"),
            ("Index",           "session, player"),
        ],
    },

    # ── LeaderboardEntry ──────────────────────────────────────────────────
    {
        "name": "LeaderboardEntry",
        "description": (
            "Snapshot del leaderboard por evento y jugador. "
            "Reconstruido por Celery tras cada batch de eventos materializados. "
            "Publicado al móvil (leaderboard_cache) y a clientes web vía Django Channels."
        ),
        "mobile_equiv": "leaderboard_cache  (tabla local WatermelonDB)",
        "fields": BASE_FIELDS + [
            {"name": "tour_event", "type": "FK → TourEvent",       "nullable": "No", "default": "—",    "description": "Evento de competición al que corresponde esta entrada.", "fk": True},
            {"name": "player",     "type": "FK → core.User",       "nullable": "No", "default": "—",    "description": "Jugador clasificado.", "fk": True},
            {"name": "session",    "type": "FK → ScoringSession",  "nullable": "Sí", "default": "NULL", "description": "Sesión activa del jugador en el evento.", "fk": True},
            {"name": "position",   "type": "PositiveIntegerField", "nullable": "Sí", "default": "NULL", "description": "Posición en el leaderboard (1 = líder)."},
            {"name": "gross",      "type": "IntegerField",          "nullable": "Sí", "default": "NULL", "description": "Total de golpes brutos acumulados."},
            {"name": "net",        "type": "IntegerField",          "nullable": "Sí", "default": "NULL", "description": "Total de golpes netos acumulados."},
            {"name": "points",     "type": "IntegerField",          "nullable": "Sí", "default": "NULL", "description": "Total de puntos Stableford acumulados."},
            {"name": "thru",       "type": "PositiveSmallIntegerField", "nullable": "Sí", "default": "NULL", "description": "Número de hoyos completados (0-18)."},
            {"name": "vs_par",     "type": "IntegerField",          "nullable": "Sí", "default": "NULL", "description": "Golpes brutos menos el par de los hoyos jugados. NULL hasta la primera materialización."},
            {"name": "status",     "type": "CharField(20)",          "nullable": "Sí", "default": "NULL", "description": "Estado del jugador en la sesión. Choices: ScoringSessionPlayerStatus."},
        ],
        "relations": [
            ("tour_event", "tour.TourEvent",          "CASCADE",   "Eliminar el evento borra todas sus entradas del leaderboard."),
            ("player",     "core.User",               "CASCADE",   "Eliminar el usuario borra sus entradas."),
            ("session",    "scoring.ScoringSession",  "SET_NULL",  "La sesión puede borrarse; la entrada permanece."),
        ],
        "indexes": [
            ("unique_together", "tour_event, player"),
            ("Index",           "tour_event, position"),
            ("ordering",        "tour_event, position"),
        ],
    },

    # ── ScoringMediaAttachment ────────────────────────────────────────────
    {
        "name": "ScoringMediaAttachment",
        "description": (
            "Archivos multimedia subidos desde el móvil, vinculados a un evento de puntuación "
            "mediante media_ref. El pipeline de subida es independiente del event log: un "
            "ScoringEvent puede referenciar un media_ref antes de que el archivo se haya subido; "
            "linked_event se rellena cuando ambos lados llegan."
        ),
        "mobile_equiv": "(sin tabla local — upload directo al servidor)",
        "fields": BASE_FIELDS + [
            {"name": "session",       "type": "FK → ScoringSession", "nullable": "No", "default": "—",    "description": "Sesión a la que pertenece el archivo.", "fk": True},
            {"name": "file",          "type": "FK → core.File",      "nullable": "No", "default": "—",    "description": "Registro de fichero en el sistema de almacenamiento.", "fk": True},
            {"name": "media_ref",     "type": "UUIDField",            "nullable": "No", "default": "—",    "description": "UUID v7 generado por el cliente. Coincide con ScoringEvent.payload.media_ref.", "unique": True},
            {"name": "linked_event",  "type": "FK → ScoringEvent",   "nullable": "Sí", "default": "NULL", "description": "Evento al que se vincula el archivo. Se rellena cuando llegan ambos lados.", "fk": True},
            {"name": "uploaded_by",   "type": "FK → core.User",      "nullable": "Sí", "default": "NULL", "description": "Usuario que subió el archivo.", "fk": True},
        ],
        "relations": [
            ("session",      "scoring.ScoringSession", "CASCADE",   "Eliminar la sesión borra sus adjuntos."),
            ("file",         "core.File",              "CASCADE",   "Eliminar el fichero borra el adjunto."),
            ("linked_event", "scoring.ScoringEvent",   "SET_NULL",  "El evento puede borrarse; el adjunto permanece."),
            ("uploaded_by",  "core.User",              "SET_NULL",  "El usuario puede borrarse sin perder el adjunto."),
        ],
        "indexes": [
            ("Index",  "session"),
            ("unique", "media_ref"),
        ],
    },
]

ENUMS = [
    ("ScoringSessionMode", "teal", [
        ("competition", "Una sesión de competición oficial."),
        ("free-play",   "Partida libre (práctica o informal)."),
    ]),
    ("ScoringSessionStatus", "teal", [
        ("in_progress", "Ronda actualmente en juego."),
        ("finished",    "Ronda completada con normalidad."),
        ("suspended",   "Ronda suspendida temporalmente."),
        ("locked",      "Tarjeta firmada y bloqueada. Sin enmiendas sin admin."),
        ("abandoned",   "Ronda abandonada antes de completarse."),
    ]),
    ("ScoringSessionPlayerStatus", "green", [
        ("not_started", "El jugador aún no se ha unido. Estado inicial."),
        ("ready",       "Dispositivo vinculado; jugador listo para empezar."),
        ("playing",     "Ronda en curso para este jugador."),
        ("finished",    "El jugador completó los 18 hoyos. Terminal."),
        ("withdrawn",   "El jugador se retiró. Terminal."),
    ]),
    ("ScoringActionType", "amber", [
        ("HOLE_SAVED",       "Golpes del hoyo guardados (shipped en mobile)."),
        ("PLAYER_READY",     "Jugador listo — dispositivo vinculado (shipped en mobile)."),
        ("ROUND_STARTED",    "Inicio de ronda (shipped en mobile)."),
        ("ROUND_FINISHED",   "Fin de ronda (shipped en mobile)."),
        ("SCORE_AMENDED",    "Corrección de un score ya guardado."),
        ("PENALTY_ADDED",    "Penalización añadida en un hoyo."),
        ("NOTE_ADDED",       "Nota de texto adjuntada a un hoyo o jugador."),
        ("MEDIA_ATTACHED",   "Archivo multimedia vinculado al evento."),
        ("SIGNATURE_ADDED",  "Firma digital de la tarjeta."),
        ("ROUND_SUSPENDED",  "Ronda suspendida (meteorología, oscuridad, etc.)."),
        ("ROUND_RESUMED",    "Ronda reanudada tras suspensión."),
        ("ADMIN_CORRECTION", "Corrección realizada por un administrador."),
        ("CONCESSION",       "Concesión de hoyo (matchplay)."),
        ("HOLE_WON",         "Hoyo ganado (matchplay)."),
        ("HOLE_HALVED",      "Hoyo empatado (matchplay)."),
    ]),
    ("ScoringMaterializationStatus", "amber", [
        ("pending",    "Evento recibido, pendiente de replay."),
        ("processed",  "Evento reproducido y MaterializedScore actualizado."),
        ("skipped",    "Evento superado por un evento posterior (ej. SCORE_AMENDED)."),
        ("failed",     "El replay falló; ver materialization_error."),
    ]),
    ("TeeColor", "purple", [
        ("yellow",         "Tees amarillos (distancia media, categoría masculina estándar)."),
        ("blue",           "Tees azules (distancia larga, categoría de élite)."),
        ("red",            "Tees rojos (distancia corta, categoría femenina estándar)."),
        ("white",          "Tees blancos (distancia larga-media)."),
        ("black",          "Tees negros (distancia máxima, torneos profesionales)."),
        ("not_applicable", "Sin categoría de tee (ej. partida libre sin recorrido configurado)."),
    ]),
]

PLAYER_TRANSITIONS = [
    ("not_started", "→ ready, → withdrawn"),
    ("ready",       "→ playing, → withdrawn"),
    ("playing",     "→ finished, → withdrawn"),
    ("finished",    "(terminal — sin transiciones)"),
    ("withdrawn",   "(terminal — sin transiciones)"),
]

# ─────────────────────────── BUILD ──────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Title page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(t, "GRANKERS BACKEND", bold=True, size_pt=26, color=P["dark_blue"])
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(s, "Modelo de Datos — scoring/models.py", bold=False, size_pt=15, color=P["green"])
    doc.add_paragraph()
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(d, "Mayo 2026  ·  Versión 1.0", size_pt=11, color=P["mid_gray"])
    doc.add_paragraph()
    doc.add_paragraph()

    # summary table
    sum_data = [
        ("ScoringSession",         "Ronda completa de un grupo",             "rounds"),
        ("ScoringSessionPlayer",   "Jugador dentro de una sesión",           "round_players"),
        ("ScoringEvent",           "Log append-only de acciones",            "action_log"),
        ("MaterializedScore",      "Score por hoyo/jugador (proyección)",    "hole_scores"),
        ("LeaderboardEntry",       "Snapshot de clasificación por evento",   "leaderboard_cache"),
        ("ScoringMediaAttachment", "Multimedia vinculada a un evento",       "— (solo servidor)"),
    ]
    sum_tbl = doc.add_table(rows=1, cols=3)
    sum_tbl.style = 'Table Grid'
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Modelo", "Descripción corta", "Tabla mobile"]):
        set_cell_bg(sum_tbl.rows[0].cells[i], HEX["dark_blue"])
        r = sum_tbl.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(*P["white"])
    for idx, (m, desc, mob) in enumerate(sum_data):
        tr = sum_tbl.add_row()
        color_key, _ = MODEL_COLOR[m]
        set_cell_bg(tr.cells[0], HEX[color_key])
        r0 = tr.cells[0].paragraphs[0].add_run(m)
        r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = RGBColor(*P["white"])
        bg = HEX["stripe_a"] if idx % 2 == 0 else HEX["stripe_b"]
        set_cell_bg(tr.cells[1], bg)
        r1 = tr.cells[1].paragraphs[0].add_run(desc)
        r1.font.size = Pt(9)
        set_cell_bg(tr.cells[2], bg)
        r2 = tr.cells[2].paragraphs[0].add_run(mob)
        r2.font.size = Pt(9); r2.font.name = "Consolas"
        r2.font.color.rgb = RGBColor(0x11, 0x55, 0x88)
    for row in sum_tbl.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(3.2)
        row.cells[2].width = Inches(1.7)

    doc.add_page_break()

    # ── BaseModel note ───────────────────────────────────────────────────────
    section_divider(doc, "BaseModel — Campos heredados por todos los modelos", "dark_blue")
    p_base = doc.add_paragraph(
        "Todos los modelos de scoring heredan de BaseModel. "
        "BaseModel añade automáticamente los campos id, uuid, created_at y updated_at. "
        "En las tablas de cada modelo estos campos se muestran con fondo gris claro."
    )
    p_base.runs[0].font.size = Pt(9)
    field_table(doc, BASE_FIELDS, HEX["dark_blue"], "E8F4FA", HEX["stripe_b"])

    # ── Models ───────────────────────────────────────────────────────────────
    for m in MODELS:
        color_key, bg_key = MODEL_COLOR[m["name"]]
        model_header(doc, m["name"], m["description"], m["mobile_equiv"])

        section_divider(doc, "Campos", color_key)
        field_table(doc, m["fields"], HEX[color_key], HEX[bg_key], HEX["stripe_b"])

        section_divider(doc, "Relaciones (Foreign Keys)", "dark_blue")
        rel_table(doc, m["relations"])

        section_divider(doc, "Índices y Constraints", "dark_blue")
        index_table(doc, m["indexes"])

        if m.get("status_transitions"):
            section_divider(doc, "Transiciones de estado — ScoringSessionPlayer.status", "green")
            note = doc.add_paragraph(
                "El dict PLAYER_STATUS_TRANSITIONS en models.py define las transiciones legales. "
                "Los estados finished y withdrawn son terminales (sin transiciones de salida)."
            )
            note.runs[0].font.size = Pt(9)
            doc.add_paragraph()
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = 'Table Grid'
            for i, h in enumerate(["Estado actual", "Puede pasar a"]):
                set_cell_bg(tbl.rows[0].cells[i], HEX["green"])
                r = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
                r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(*P["white"])
            for idx, (state, targets) in enumerate(PLAYER_TRANSITIONS):
                tr = tbl.add_row()
                bg = HEX["light_green"] if idx % 2 == 0 else HEX["stripe_b"]
                set_cell_bg(tr.cells[0], bg); set_cell_bg(tr.cells[1], bg)
                r0 = tr.cells[0].paragraphs[0].add_run(state)
                r0.font.size = Pt(8); r0.font.name = "Consolas"
                r0.font.color.rgb = RGBColor(0x11, 0x55, 0x44)
                r1 = tr.cells[1].paragraphs[0].add_run(targets)
                r1.font.size = Pt(8)
            for row in tbl.rows:
                row.cells[0].width = Inches(1.8); row.cells[1].width = Inches(5.3)
            doc.add_paragraph()

    # ── Enums ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    section_divider(doc, "Enumeraciones (core/constants.py)", "dark_blue")
    p_enum = doc.add_paragraph(
        "Los campos CharField con choices utilizan las siguientes enumeraciones "
        "definidas en core/constants.py."
    )
    p_enum.runs[0].font.size = Pt(9)
    doc.add_paragraph()
    for (name, color, vals) in ENUMS:
        enum_box(doc, name, vals, color)

    out = r"c:\developments\GrankersMobile\docs\Scoring_Models.docx"
    doc.save(out)
    print(f"Guardado: {out}")

if __name__ == "__main__":
    build()
