"""
Genera GrankersMobile_Pantallas.docx con la documentación de todas las pantallas.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ───────────────────────────── helpers ──────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_placeholder_img(doc, label="[Captura de pantalla]"):
    """Añade una caja de placeholder con borde para la imagen."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(3.5)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2D9B4E')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    set_cell_bg(cell, 'F0FFF4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(label)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2D, 0x9B, 0x4E)
    run.font.italic = True
    doc.add_paragraph()

def add_table_section(doc, title, rows):
    """rows = list of (campo, descripcion) tuples."""
    if not rows:
        return
    p = doc.add_paragraph()
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header
    hdr = tbl.rows[0].cells
    set_cell_bg(hdr[0], '1A5676')
    set_cell_bg(hdr[1], '1A5676')
    for idx, txt in enumerate(['Campo / Elemento', 'Descripción']):
        p2 = hdr[idx].paragraphs[0]
        run2 = p2.add_run(txt)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run2.font.size = Pt(9)

    # data rows
    for i, (campo, desc) in enumerate(rows):
        row = tbl.add_row()
        bg = 'F5FBFF' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        r0 = row.cells[0].paragraphs[0].add_run(campo)
        r0.font.size = Pt(9)
        r0.bold = True
        r1 = row.cells[1].paragraphs[0].add_run(desc)
        r1.font.size = Pt(9)

    # column widths
    for row in tbl.rows:
        row.cells[0].width = Inches(1.9)
        row.cells[1].width = Inches(4.1)

    doc.add_paragraph()

def screen_header(doc, number, title, route, description):
    doc.add_page_break()
    h = doc.add_heading('', level=2)
    h.clear()
    run = h.add_run(f"{number}. {title}")
    run.font.color.rgb = RGBColor(0x0F, 0x3E, 0x5A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Ruta: ")
    r.bold = True
    r.font.size = Pt(9)
    r2 = p.add_run(route)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if description:
        pd = doc.add_paragraph(description)
        pd.paragraph_format.space_after = Pt(6)
        pd.runs[0].font.size = Pt(9)

# ───────────────────────────── DATA ─────────────────────────────────────────

SCREENS = [
    # ---- AUTENTICACIÓN ----
    {
        "section": "AUTENTICACIÓN",
        "number": "1",
        "title": "Pantalla de Bienvenida",
        "route": "/  (app/index.tsx)",
        "description": (
            "Primera pantalla que ve el usuario al abrir la app. Presenta el logo de Grankers "
            "y las opciones de acceso. Si el usuario ya está autenticado, muestra directamente el menú principal."
        ),
        "campos": [
            ("Logo Grankers", "Logotipo de la aplicación centrado en pantalla."),
            ("Tagline", '"Golfers only beyond this point" — eslogan de la app.'),
            ("Nombre de usuario", "Se muestra si el usuario ya está autenticado (nombre + apellido desde el perfil)."),
            ("Estado de conexión", "Indicador en el header: Online (verde), Reconectando (ámbar pulsante), Sin conexión (rojo)."),
        ],
        "botones": [
            ("Iniciar sesión", "Navega a la pantalla de login (/player-area/login)."),
            ("Crear cuenta", "Navega al formulario de registro (/player-area/register)."),
            ("Continuar con Google", "Inicia el flujo OAuth de Google; tras autenticarse, navega al menú principal."),
            ("Área del Jugador", "Visible solo si autenticado. Navega a /player-area (perfil, eventos, licencias)."),
            ("Competición", "Comprueba si hay competición activa; navega a /competition/code-entry."),
            ("Partida Libre", "Navega al selector de campo para partida libre (/free-play/select-course)."),
            ("Cerrar sesión", "Visible solo si autenticado. Cierra la sesión y vuelve al estado no autenticado."),
        ],
    },
    {
        "section": None,
        "number": "2",
        "title": "Inicio de Sesión",
        "route": "/player-area/login  (app/player-area/login.tsx)",
        "description": (
            "Pantalla de acceso mediante magic link por email o con cuenta de Google. "
            "Tras enviar el enlace, la pantalla confirma el envío."
        ),
        "campos": [
            ("Email", "Campo de texto para introducir la dirección de correo electrónico (teclado tipo email)."),
            ("Confirmación post-envío", "Después de enviar, muestra el email al que se ha enviado el enlace de acceso."),
        ],
        "botones": [
            ("Iniciar sesión con Google", "Inicia el flujo OAuth de Google directamente desde esta pantalla."),
            ("Enviar enlace de acceso", "Envía un magic link al email introducido. La pantalla pasa a modo de confirmación."),
            ("¿No tienes una cuenta? Regístrese aquí", "Navega a /player-area/register."),
            ("Volver", "Regresa a la pantalla anterior."),
        ],
    },
    {
        "section": None,
        "number": "3",
        "title": "Registro de Usuario",
        "route": "/player-area/register  (app/player-area/register.tsx)",
        "description": (
            "Formulario para crear una nueva cuenta. Requiere nombre, apellidos, email y país. "
            "Incluye aceptación de términos. También permite registro con Google."
        ),
        "campos": [
            ("Nombre", "Texto libre, autoCapitalize: words. Obligatorio."),
            ("Apellidos", "Texto libre, autoCapitalize: words. Obligatorio."),
            ("Correo electrónico", "Email, teclado tipo email-address. Obligatorio."),
            ("País de residencia", "Selector modal con 45 países disponibles. Obligatorio."),
            ("Acepto términos y política de privacidad", "Checkbox obligatorio para poder crear la cuenta."),
        ],
        "botones": [
            ("Regístrese con Google", "Crea cuenta vinculada a la cuenta de Google del usuario."),
            ("Crear cuenta", "Envía el formulario y crea la cuenta; muestra confirmación para verificar email."),
            ("Términos de uso", "Enlace que abre /player-area/terms."),
            ("Política de privacidad", "Enlace que abre /player-area/privacy."),
            ("¿Ya tienes una cuenta? Inicia sesión aquí", "Navega a /player-area/login."),
        ],
    },

    # ---- ÁREA DEL JUGADOR ----
    {
        "section": "ÁREA DEL JUGADOR",
        "number": "4",
        "title": "Área del Jugador — Inicio",
        "route": "/player-area  (app/player-area/index.tsx)",
        "description": (
            "Panel principal del perfil del jugador. Muestra datos del perfil, estadísticas de "
            "actividad y los próximos eventos registrados."
        ),
        "campos": [
            ("Tarjeta de perfil", "Avatar (inicial si no hay foto), nombre completo y email del usuario."),
            ("Eventos (contador)", "Número total de eventos en los que el jugador está inscrito."),
            ("Torneos (contador)", "Número de torneos distintos en los que ha participado."),
            ("Clubes (contador)", "Número de clubes de golf con los que ha tenido actividad."),
            ("Próximos eventos", "Lista de hasta 5 eventos futuros: nombre del evento, torneo, fecha, club, estado y precio."),
            ("Estado del evento", "Chip de color: Confirmado (verde), Pendiente (ámbar), Cancelado (rojo)."),
            ("Estado de pago", "Indica si está Pagado, Pago pendiente o si es Gratuito."),
        ],
        "botones": [
            ("Tarjeta de perfil [tap]", "Navega a /player-area/profile para editar datos personales y licencias."),
            ("Competiciones", "Navega a /player-area/competitions (historial completo de inscripciones)."),
            ("Tarjetas de puntuación", "Navega a /player-area/scorecard (historial de rondas jugadas)."),
            ("Licencias y hándicap", "Navega a /player-area/profile (tab Golf)."),
            ("Mi perfil", "Navega a /player-area/profile (tab Personal)."),
            ("Ajustes", "Navega a /player-area/settings."),
            ("Cerrar sesión", "Cierra la sesión activa y regresa a la pantalla de bienvenida."),
        ],
    },
    {
        "section": None,
        "number": "5",
        "title": "Mi Perfil",
        "route": "/player-area/profile  (app/player-area/profile.tsx)",
        "description": (
            "Pantalla con dos pestañas: Personal (datos del jugador) y Golf (licencias federativas y hándicap)."
        ),
        "campos": [
            ("TAB PERSONAL — Nombre", "Campo editable con el nombre del jugador."),
            ("TAB PERSONAL — Apellidos", "Campo editable con los apellidos."),
            ("TAB PERSONAL — Teléfono", "Campo editable; teclado phone-pad. Placeholder: +34 600 000 000."),
            ("TAB PERSONAL — Fecha de nacimiento", "Campo editable en formato AAAA-MM-DD."),
            ("TAB PERSONAL — Género", "Selector modal: Masculino, Femenino, Otro, No especificado, Prefiero no decirlo."),
            ("TAB PERSONAL — Idioma", "Selector modal: Español, English, Català, Français, Deutsch, Italiano, Português."),
            ("TAB PERSONAL — Email", "Solo lectura; no editable desde la app."),
            ("TAB PERSONAL — Miembro desde", "Año de alta en la plataforma. Solo lectura."),
            ("TAB GOLF — Federación", "Selector modal con la lista de federaciones de golf disponibles."),
            ("TAB GOLF — Número de licencia", "Texto libre, autoCapitalize: characters. Ej. 12345678A."),
            ("TAB GOLF — Hándicap", "Numérico decimal opcional. Ej. 18.4."),
        ],
        "botones": [
            ("Guardar cambios (Personal)", "Persiste los datos personales editados en el servidor."),
            ("Añadir licencia", "Despliega el formulario inline para introducir una nueva licencia de golf."),
            ("Editar licencia", "Habilita la edición de una licencia existente."),
            ("Eliminar licencia", "Elimina la licencia seleccionada (pide confirmación)."),
            ("Guardar licencia", "Guarda la licencia nueva o editada."),
            ("Cancelar licencia", "Cancela la edición sin guardar."),
        ],
    },
    {
        "section": None,
        "number": "6",
        "title": "Ajustes",
        "route": "/player-area/settings  (app/player-area/settings.tsx)",
        "description": (
            "Pantalla de configuración de la app: notificaciones push, apariencia, privacidad y zona de peligro."
        ),
        "campos": [
            ("Competiciones (toggle)", "Activa/desactiva notificaciones push sobre cambios en inscripciones y resultados."),
            ("Resultados (toggle)", "Notificaciones cuando se publiquen puntuaciones."),
            ("Torneos disponibles (toggle)", "Notificaciones sobre nuevos eventos en clubes."),
            ("Comunicaciones (toggle)", "Novedades y promociones de Grankers."),
            ("Modo oscuro (toggle)", "Activa el tema oscuro en la interfaz de la app."),
            ("Versión de la app", "Texto informativo: «Grankers Mobile · v1.0». Solo lectura."),
        ],
        "botones": [
            ("Idioma", "Redirige a /player-area/profile para cambiar el idioma desde el perfil."),
            ("Política de privacidad", "Abre /player-area/privacy con el texto legal completo."),
            ("Términos de uso", "Abre /player-area/terms con los términos de uso."),
            ("Mis datos", "Abre el cliente de correo para contactar con soporte@grankers.com."),
            ("Eliminar cuenta", "Inicia el proceso de eliminación de cuenta (requiere contacto con soporte)."),
        ],
    },
    {
        "section": None,
        "number": "7",
        "title": "Mis Competiciones",
        "route": "/player-area/competitions  (app/player-area/competitions.tsx)",
        "description": (
            "Listado completo de todas las inscripciones del jugador, organizadas en próximos y "
            "historial, con filtros por estado."
        ),
        "campos": [
            ("Filtro de estado", "Chips seleccionables: Todos | Confirmados | Pendientes."),
            ("Nombre del evento", "Nombre de la competición/prueba."),
            ("Torneo + fecha", "Nombre del torneo y fecha de celebración."),
            ("Club de golf", "Campo de golf donde se celebra."),
            ("Estado", "Chip de color: Confirmado, Pendiente o Cancelado."),
            ("Categoría de cuota", "Categoría de inscripción con su precio."),
            ("Estado de pago", "Pagado / Pago pendiente / Gratuito."),
        ],
        "botones": [
            ("Filtro Todos / Confirmados / Pendientes", "Filtra la lista por estado de inscripción."),
            ("[Tarjeta de evento]", "Muestra los detalles completos de esa competición."),
        ],
    },
    {
        "section": None,
        "number": "8",
        "title": "Tarjetas de Puntuación (historial)",
        "route": "/player-area/scorecard  (app/player-area/scorecard.tsx)",
        "description": (
            "Historial de rondas jugadas por el jugador, con los mismos filtros que la pantalla de competiciones."
        ),
        "campos": [
            ("Filtro de estado", "Chips seleccionables: Todos | Confirmados | Pendientes."),
            ("Nombre del evento", "Evento al que corresponde la tarjeta."),
            ("Torneo + fecha", "Torneo y fecha de la ronda."),
            ("Club de golf", "Campo donde se jugó."),
            ("Estado con punto de color", "Indicador visual del estado de la competición."),
            ("Categoría y precio", "Categoría de inscripción y precio pagado."),
        ],
        "botones": [
            ("[Tarjeta de evento]", "Abre los detalles de esa ronda / tarjeta de puntuación."),
        ],
    },
    {
        "section": None,
        "number": "9",
        "title": "Términos de Uso",
        "route": "/player-area/terms  (app/player-area/terms.tsx)",
        "description": "Texto legal desplazable con los términos de uso de la aplicación. Actualizado febrero 2026.",
        "campos": [
            ("Contenido legal", "Secciones: Aceptación, Descripción del servicio, Registro, Uso aceptable, Propiedad intelectual, Limitación de responsabilidad, Modificaciones, Contacto."),
        ],
        "botones": [
            ("Volver (header)", "Regresa a la pantalla anterior."),
        ],
    },
    {
        "section": None,
        "number": "10",
        "title": "Política de Privacidad",
        "route": "/player-area/privacy  (app/player-area/privacy.tsx)",
        "description": "Texto legal desplazable con la política de privacidad. Actualizado febrero 2026.",
        "campos": [
            ("Contenido legal", "Secciones: Información recopilada, Uso de información, Compartir información, Almacenamiento, Derechos del usuario, Cookies, Cambios en la política, Contacto."),
        ],
        "botones": [
            ("Volver (header)", "Regresa a la pantalla anterior."),
        ],
    },

    # ---- COMPETICIÓN ----
    {
        "section": "MODO COMPETICIÓN",
        "number": "11",
        "title": "Entrada al Grupo de Competición",
        "route": "/competition/code-entry  (app/competition/code-entry.tsx)",
        "description": (
            "Pantalla inteligente con tres estados según si el jugador tiene o no una competición "
            "registrada para hoy. "
            "(A) Competición detectada automáticamente: muestra detalles de la prueba y solicita el código de grupo. "
            "(B) Grupo cargado: muestra todos los jugadores del grupo, marcador asignado y botón Listo. "
            "(C) Sin competición hoy: muestra mensaje explicativo y permite entrada manual del código."
        ),
        "campos": [
            ("ESTADO A — Nombre del torneo", "Nombre del torneo al que pertenece la competición de hoy."),
            ("ESTADO A — Nombre de la prueba", "Nombre de la prueba/evento dentro del torneo."),
            ("ESTADO A — Fecha", "Fecha de celebración de la competición."),
            ("ESTADO A — Tu hándicap", "Hándicap de la licencia activa del jugador (valor de referencia; playing handicap definitivo viene del backend)."),
            ("ESTADO A — Código de grupo", "Campo de texto (autoCapitalize: characters, máx. 12 caracteres). Código impreso en la tarjeta de puntuación en papel."),
            ("ESTADO B — Lista de jugadores", "Filas por cada jugador del grupo: nombre, hándicap y rol (Marcas a: / Tu marcador)."),
            ("ESTADO B — Tu fila (destacada)", "La fila del jugador autenticado aparece con fondo diferenciado."),
            ("ESTADO C — Mensaje informativo", "«Su usuario no aparece como inscrito en ninguna competición para el día de hoy...»"),
            ("ESTADO C — Código de grupo", "Campo de texto para entrada manual del código (mismo campo que estado A)."),
        ],
        "botones": [
            ("→ (icono envío)", "Envía el código de grupo al servidor y carga los jugadores del grupo."),
            ("Listo (estado B)", "Confirma el grupo y navega a /competition/select-player para vincular el dispositivo."),
        ],
    },
    {
        "section": None,
        "number": "12",
        "title": "Seleccionar Jugador (Competición)",
        "route": "/competition/select-player  (app/competition/select-player.tsx)",
        "description": (
            "El jugador identifica en qué dispositivo está: selecciona su nombre de la lista. "
            "El sistema vincula el dispositivo a ese jugador en el servidor."
        ),
        "campos": [
            ("Título", "«¿Quién eres?» — instrucción para que el usuario seleccione su nombre."),
            ("Lista de jugadores", "Solo muestra los jugadores que aún no tienen dispositivo vinculado. Muestra nombre y hándicap."),
        ],
        "botones": [
            ("[Tarjeta de jugador]", "Vincula el dispositivo a ese jugador (llamada al endpoint link-device) y navega a /competition/waiting-players."),
        ],
    },
    {
        "section": None,
        "number": "13",
        "title": "Esperando Jugadores",
        "route": "/competition/waiting-players  (app/competition/waiting-players.tsx)",
        "description": (
            "Sala de espera hasta que todos los miembros del grupo hayan vinculado su dispositivo. "
            "Muestra en tiempo real el estado de conexión de cada jugador."
        ),
        "campos": [
            ("Contador de conexiones", "«X / Y jugadores conectados» — actualizado en tiempo real."),
            ("Lista de jugadores", "Nombre de cada jugador + estado: ✓ Conectado / ⟳ Esperando / WiFi-off Offline."),
            ("Indicador (Tú)", "La fila del propio jugador está marcada con «(Tú)»."),
        ],
        "botones": [
            ("Checkbox Offline", "Marca a un jugador como offline para que no bloquee el inicio."),
            ("Todos los jugadores (modal)", "El dispositivo puntuará a todos los jugadores del grupo."),
            ("Solo el que me corresponde (modal)", "El dispositivo puntuará solo al jugador asignado (modo parcial)."),
        ],
    },
    {
        "section": None,
        "number": "14",
        "title": "Puntuación — Competición",
        "route": "/competition/scoring  (app/competition/scoring.tsx)",
        "description": (
            "Pantalla principal de entrada de golpes durante la ronda de competición. "
            "Muestra un hoyo a la vez con los controles de puntuación de cada jugador del grupo."
        ),
        "campos": [
            ("Header — Nombre de competición/evento", "Nombre de la prueba en la que se está jugando."),
            ("Header — Hoyo actual", "Número del hoyo (1-18) que se está puntuando ahora."),
            ("Header — Par del hoyo", "Par oficial del hoyo según el recorrido."),
            ("Header — HCP del hoyo", "Índice de hándicap del hoyo (1 = más difícil, 18 = más fácil)."),
            ("Tarjeta de jugador — Nombre", "Nombre del jugador; resaltado si es el propio dispositivo o el marcador asignado."),
            ("Tarjeta de jugador — Tag «Tú»", "Etiqueta azul indicando que es el jugador del dispositivo."),
            ("Tarjeta de jugador — Tag «Marcando»", "Etiqueta dorada indicando que es el marcador del dispositivo."),
            ("Tarjeta de jugador — Score", "Número de golpes registrado para ese hoyo. Editable con los botones ± ."),
            ("Tarjeta de jugador — Diferencial vs par", "Resultado respecto al par: −1 (birdie), E (par), +1 (bogey)…"),
            ("Tarjeta de jugador — Puntos Stableford", "Puntos Stableford calculados automáticamente según el hándicap."),
        ],
        "botones": [
            ("− (menos)", "Decrementa un golpe para el jugador correspondiente."),
            ("+ (más)", "Incrementa un golpe para el jugador correspondiente."),
            ("Atrás", "Regresa al hoyo anterior para corrección."),
            ("Guardar / Siguiente", "Guarda el hoyo actual y avanza al siguiente (en el hoyo 18 navega a /competition/comprobacion)."),
            ("Editar", "Habilita el modo edición para modificar scores ya guardados."),
            ("Ranking (trofeo footer)", "Abre /game/leaderboard con la clasificación del grupo en tiempo real."),
            ("Tarjeta (footer)", "Abre /game/scorecard con la tarjeta de puntuación completa."),
            ("Abandonar (rojo footer)", "Muestra confirmación para salir de la competición y regresa a /."),
        ],
    },
    {
        "section": None,
        "number": "15",
        "title": "Comprobación de Resultado",
        "route": "/competition/comprobacion  (app/competition/comprobacion.tsx)",
        "description": (
            "Pantalla de cierre de ronda donde el jugador y su marcador verifican y firman "
            "la tarjeta antes de enviarla al servidor."
        ),
        "campos": [
            ("Golpes registrados por el marcador", "Nombre del marcador, total de golpes y diferencial vs par."),
            ("Golpes apuntados por mí", "Mi nombre, total de golpes y diferencial vs par."),
            ("Estado «Esperando»", "Mensaje que aparece si el marcador aún no ha completado su tarjeta."),
        ],
        "botones": [
            ("Revisar tarjeta (ojo)", "Abre un modal con la tabla hoyo a hoyo del marcador para revisión detallada."),
            ("Firmar tarjeta (check)", "Firma digitalmente la tarjeta tras confirmación; envía al servidor y navega a /."),
            ("Modal — Editar hoyo (±)", "Permite corregir un score individual dentro del modal de revisión."),
            ("Modal — Guardar cambios", "Guarda los cambios realizados en el modal."),
        ],
    },

    # ---- PARTIDA LIBRE ----
    {
        "section": "PARTIDA LIBRE",
        "number": "16",
        "title": "Seleccionar Campo y Recorrido",
        "route": "/free-play/select-course  (app/free-play/select-course.tsx)",
        "description": (
            "Primer paso de la partida libre: elegir el campo de golf y el recorrido. "
            "Si ya hay partidas activas en ese recorrido, ofrece unirse a una existente."
        ),
        "campos": [
            ("Buscador de campo", "Campo de texto con filtrado en tiempo real sobre la lista de campos disponibles."),
            ("Lista de campos de golf", "Lista desplazable de todos los campos disponibles en el catálogo."),
            ("Lista de recorridos", "Una vez seleccionado el campo, muestra sus recorridos disponibles."),
            ("Modal — Partidas activas", "Lista de sesiones en curso en ese recorrido (nombre, número de jugadores)."),
        ],
        "botones": [
            ("Continuar", "Verifica partidas activas en el recorrido y navega según el resultado."),
            ("Crear Partida Nueva (modal)", "Ignora las partidas existentes y navega a /free-play/create-game."),
            ("[Partida existente] (modal)", "Se une a esa sesión y navega directamente a /free-play/select-device-player."),
            ("X cerrar modal", "Cierra el modal de partidas activas."),
        ],
    },
    {
        "section": None,
        "number": "17",
        "title": "Crear Partida",
        "route": "/free-play/create-game  (app/free-play/create-game.tsx)",
        "description": "Configuración de los parámetros básicos de la nueva partida libre.",
        "campos": [
            ("Campo de golf", "Solo lectura — campo seleccionado en la pantalla anterior."),
            ("Recorrido", "Solo lectura — recorrido seleccionado."),
            ("Nombre de la partida", "Texto libre, máx. 80 caracteres. Prerrelleno con el nombre del campo + «_partida»."),
            ("Jugadores por grupo", "Contador con botones −/+ (rango 1–4). Muestra «X jugadores / jugador»."),
            ("Tipo de partida", "Selector: Pública (icono globo) o Privada (icono candado)."),
            ("Contraseña", "Solo visible si tipo = Privada. Texto oculto, máx. 20 caracteres."),
        ],
        "botones": [
            ("Añadir Jugadores →", "Crea la partida en el servidor y navega a /free-play/setup."),
            ("Cancelar", "Descarta la partida y regresa a la pantalla anterior."),
        ],
    },
    {
        "section": None,
        "number": "18",
        "title": "Configurar Jugadores",
        "route": "/free-play/setup  (app/free-play/setup.tsx)",
        "description": (
            "Introducción manual de los datos de cada jugador del grupo o búsqueda por licencia. "
            "Se muestra una PlayerCard por cada jugador según el número configurado."
        ),
        "campos": [
            ("Nombre (por jugador)", "Campo editable con el nombre del jugador."),
            ("Apellido (por jugador)", "Campo editable con el apellido."),
            ("Hándicap (por jugador)", "Campo numérico editable con el hándicap de juego."),
        ],
        "botones": [
            ("Buscar Licencia (por jugador)", "Navega a /free-play/search-license para rellenar los datos desde la base de jugadores federados."),
            ("Siguiente →", "Guarda los datos de los jugadores y navega a /free-play/select-device-player."),
            ("Cancelar", "Descarta la partida y regresa a /."),
        ],
    },
    {
        "section": None,
        "number": "19",
        "title": "Buscar por Licencia",
        "route": "/free-play/search-license  (app/free-play/search-license.tsx)",
        "description": (
            "Búsqueda de jugadores federados por número de licencia, nombre o apellido. "
            "Al seleccionar un resultado, rellena automáticamente los campos del jugador en /free-play/setup."
        ),
        "campos": [
            ("Licencia", "Número de licencia federativa. autoCapitalize: characters."),
            ("Nombre", "Nombre del jugador a buscar."),
            ("Apellido", "Apellido del jugador a buscar."),
            ("Resultados encontrados", "Contador de resultados y lista de tarjetas de jugadores encontrados."),
        ],
        "botones": [
            ("Buscar (icono lupa)", "Ejecuta la búsqueda en la base de datos de jugadores federados."),
            ("[Tarjeta de resultado]", "Selecciona el jugador y transfiere sus datos a /free-play/setup."),
        ],
    },
    {
        "section": None,
        "number": "20",
        "title": "Seleccionar Jugador del Dispositivo (Partida Libre)",
        "route": "/free-play/select-device-player  (app/free-play/select-device-player.tsx)",
        "description": (
            "Cada dispositivo elige qué jugador del grupo representa. "
            "Paso final antes de iniciar la puntuación."
        ),
        "campos": [
            ("Título", "«¿Quién eres?» — instrucción para identificar al jugador del dispositivo."),
            ("Tarjetas de jugadores", "Una tarjeta por cada jugador configurado: nombre y hándicap."),
        ],
        "botones": [
            ("[Tarjeta de jugador]", "Asigna este dispositivo al jugador seleccionado e inicia la sesión de puntuación en /game/scoring."),
        ],
    },

    # ---- JUEGO ----
    {
        "section": "PANTALLAS DE JUEGO (COMPARTIDAS)",
        "number": "21",
        "title": "Puntuación — Partida Libre",
        "route": "/game/scoring  (app/game/scoring.tsx)",
        "description": (
            "Pantalla principal de entrada de golpes para partidas libres. "
            "Misma estructura que la pantalla de puntuación de competición."
        ),
        "campos": [
            ("Header — Nombre de partida/competición", "Nombre de la partida o competición activa."),
            ("Header — Hoyo actual / Par / HCP", "Información del hoyo que se está jugando."),
            ("Tarjeta de jugador — Score", "Golpes del hoyo, editable con controles ±."),
            ("Tarjeta de jugador — Diferencial vs par", "Resultado respecto al par del hoyo."),
            ("Tarjeta de jugador — Puntos Stableford", "Puntos Stableford calculados automáticamente."),
        ],
        "botones": [
            ("− / +", "Decrementa o incrementa un golpe para el jugador."),
            ("Atrás", "Regresa al hoyo anterior."),
            ("Guardar / Siguiente", "Guarda el hoyo y avanza. En el hoyo 18 navega a /game/complete."),
            ("Editar", "Habilita la edición de scores ya guardados."),
            ("Ranking (trofeo)", "Abre /game/leaderboard."),
            ("Tarjeta (crédito)", "Abre /game/scorecard."),
            ("Abandonar (rojo)", "Abandona la partida con confirmación y regresa a /."),
        ],
    },
    {
        "section": None,
        "number": "22",
        "title": "Tarjeta de Puntuación",
        "route": "/game/scorecard  (app/game/scorecard.tsx)",
        "description": (
            "Vista completa de la tarjeta de puntuación con tabla de 18 hoyos. "
            "Dos modos: vista general (tabla) y vista de hoyo individual."
        ),
        "campos": [
            ("Nombre del jugador", "Jugador al que pertenece la tarjeta."),
            ("Resumen", "Total de golpes | Diferencial vs par (+/−) | Hoyos completados."),
            ("Tabla de hoyos", "18 filas: Hoyo | Par | Score. Colores: birdie (azul), par (amarillo), bogey (naranja), doble bogey+ (rojo)."),
            ("Vista de hoyo — Score", "Número grande de golpes para ese hoyo, modificable con ±."),
            ("Vista de hoyo — Descripción", "Texto descriptivo: «Par», «X bajo par», «X sobre par»."),
        ],
        "botones": [
            ("[Hoyo en la tabla]", "Navega a la vista de hoyo individual."),
            ("Volver a tarjeta", "Regresa a la vista de tabla general."),
            ("Anterior / Siguiente", "Navega entre hoyos en la vista individual."),
            ("Editar / Guardar", "Alterna entre modo edición y modo lectura; guarda los cambios."),
        ],
    },
    {
        "section": None,
        "number": "23",
        "title": "Clasificación (Leaderboard)",
        "route": "/game/leaderboard  (app/game/leaderboard.tsx)",
        "description": (
            "Clasificación en tiempo real del grupo durante la ronda. "
            "Se actualiza por WebSocket (o por polling REST si el WS cae)."
        ),
        "campos": [
            ("Nombre del evento / competición", "Encabezado con el nombre de la partida o prueba."),
            ("Posición", "1/2/3 con iconos de trofeo/medalla; resto con número."),
            ("Nombre del jugador", "Nombre completo. Resaltado si es el líder actual."),
            ("Hoyos completados", "«X de 18» — progreso de la ronda."),
            ("Score diferencial", "Resultado acumulado vs par: −2, E, +3…"),
            ("Score total", "Total de golpes acumulados."),
        ],
        "botones": [
            ("[Jugador]", "Abre /game/scorecard con la tarjeta detallada de ese jugador."),
            ("Volver", "Regresa a la pantalla de puntuación activa."),
        ],
    },
    {
        "section": None,
        "number": "24",
        "title": "Revisión de Hoyos",
        "route": "/game/review  (app/game/review.tsx)",
        "description": (
            "Cuadrícula de 18 hoyos para seleccionar rápidamente qué hoyo revisar o editar."
        ),
        "campos": [
            ("Instrucción", "«Selecciona un hoyo» / «Elige el hoyo que quieres revisar o editar»."),
            ("Grid de hoyos", "18 botones numerados del 1 al 18."),
        ],
        "botones": [
            ("[Número de hoyo]", "Abre /game/scorecard en la vista de ese hoyo concreto."),
        ],
    },
    {
        "section": None,
        "number": "25",
        "title": "Ronda Completada",
        "route": "/game/complete  (app/game/complete.tsx)",
        "description": (
            "Pantalla final que aparece cuando se guarda el último hoyo. "
            "Muestra la clasificación definitiva y el estado de sincronización con el servidor."
        ),
        "campos": [
            ("Título «¡Tarjeta Firmada!»", "Encabezado con icono de check verde."),
            ("Estado de sincronización", "Icono Cloud / CloudOff + mensaje sobre si los datos se han enviado al servidor."),
            ("Tabla de clasificación final", "Posición, nombre y score total de cada jugador. Score en rojo si + (sobre par), verde si − (bajo par)."),
        ],
        "botones": [
            ("Terminar", "Resetea el estado de la partida y navega a / (pantalla de bienvenida)."),
        ],
    },
]

# ───────────────────────────── BUILD DOC ────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("GRANKERS MOBILE")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x0F, 0x3E, 0x5A)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Documentación de Pantallas y Funcionalidad")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x2D, 0x9B, 0x4E)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run("Mayo 2026  ·  Versión 1.0")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note_p.add_run(
        "📱  Las capturas de pantalla deben adjuntarse en los recuadros verdes\n"
        "    tomando screenshots de la app en ejecución."
    )
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    note_run.italic = True

    doc.add_page_break()

    # ── Intro ────────────────────────────────────────────────────────────────
    doc.add_heading("Introducción", level=1)
    intro = doc.add_paragraph(
        "Este documento describe todas las pantallas de la aplicación móvil Grankers Mobile, "
        "desarrollada con Expo y React Native. Para cada pantalla se detalla: la ruta en la app, "
        "una descripción funcional, todos los campos y la información que muestran, y cada botón "
        "con la acción que realiza.\n\n"
        "La app tiene 25 pantallas principales organizadas en cinco bloques:\n"
        "  1.  Autenticación (pantallas 1–3)\n"
        "  2.  Área del Jugador (pantallas 4–10)\n"
        "  3.  Modo Competición (pantallas 11–15)\n"
        "  4.  Partida Libre (pantallas 16–20)\n"
        "  5.  Pantallas de Juego compartidas (pantallas 21–25)"
    )
    intro.runs[0].font.size = Pt(10)

    # ── Screens ──────────────────────────────────────────────────────────────
    current_section = None

    for screen in SCREENS:
        # Section heading
        if screen["section"] and screen["section"] != current_section:
            current_section = screen["section"]
            doc.add_page_break()
            h = doc.add_heading(current_section, level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x2D, 0x9B, 0x4E)

        # Screen heading + meta
        screen_header(doc, screen["number"], screen["title"], screen["route"], screen["description"])

        # Placeholder image
        add_placeholder_img(doc, f"[Captura de pantalla — {screen['title']}]")

        # Tables
        if screen.get("campos"):
            add_table_section(doc, "CAMPOS E INFORMACIÓN MOSTRADA", [
                (c[0], c[1]) for c in screen["campos"]
            ])
        if screen.get("botones"):
            add_table_section(doc, "BOTONES Y ACCIONES", [
                (b[0], b[1]) for b in screen["botones"]
            ])

    # ── Footer note ──────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Flujo de Navegación — Resumen", level=1)
    nav_text = doc.add_paragraph(
        "Inicio  →  Login / Registro  →  Área del Jugador\n"
        "Inicio  →  Competición  →  Entrada código grupo  →  Seleccionar jugador  →  Sala de espera  →  Puntuación  →  Comprobación / Firma\n"
        "Inicio  →  Partida Libre  →  Seleccionar campo  →  Crear partida  →  Configurar jugadores  →  Seleccionar dispositivo  →  Puntuación  →  Completado\n\n"
        "Las pantallas de Clasificación (Leaderboard), Tarjeta de puntuación (Scorecard) y Revisión de hoyos "
        "son accesibles desde la pantalla de Puntuación tanto en Competición como en Partida Libre."
    )
    nav_text.runs[0].font.size = Pt(10)

    out = r"c:\developments\GrankersMobile\docs\GrankersMobile_Pantallas.docx"
    doc.save(out)
    print(f"Guardado: {out}")

if __name__ == "__main__":
    build()
